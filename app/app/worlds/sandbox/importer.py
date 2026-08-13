from __future__ import annotations

import json
from pathlib import Path

from app.agents.script_decomposition import ScriptDecompositionAgent
from app.agents.world_builder import WorldBuilderAgent, generate_world_config
from app.worlds.sandbox.models import SandboxWorldConfig, ScriptDecompositionRequest, WorldGenerateRequest


MAX_IMPORT_CHARS = 20000
TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".csv", ".json", ".html", ".htm", ".rtf"}


async def import_world_from_document(
    filename: str,
    content: bytes,
    player_name: str = "",
    world_name: str = "",
    use_ai: bool = True,
) -> SandboxWorldConfig:
    text = extract_document_text(filename, content)
    if _looks_like_script_case(text):
        return ScriptDecompositionAgent().build(
            ScriptDecompositionRequest(
                title=world_name,
                player_name=player_name or "侦探",
                source_text=text,
            )
        ).world
    request = WorldGenerateRequest(
        template="document_import",
        theme=_document_theme(filename, text),
        player_name=player_name or "主角",
        world_name=world_name,
    )
    fallback = _document_fallback_world(request, filename, text)
    if not use_ai:
        return fallback
    return await WorldBuilderAgent().generate(request, fallback=fallback)


def extract_document_text(filename: str, content: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix in TEXT_SUFFIXES:
        text = _decode_text(content)
    elif suffix == ".pdf":
        text = _extract_pdf_text(content)
    elif suffix == ".docx":
        text = _extract_docx_text(content)
    else:
        text = _decode_text(content)

    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not text:
        raise ValueError("没有从文档中读取到有效文本。")
    return text[:MAX_IMPORT_CHARS]


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_pdf_text(content: bytes) -> str:
    try:
        from io import BytesIO

        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("当前环境缺少 pypdf，无法读取 PDF。请安装 pypdf。") from exc

    reader = PdfReader(BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(content: bytes) -> str:
    try:
        from io import BytesIO

        import docx
    except ImportError as exc:
        raise ValueError("当前环境缺少 python-docx，无法读取 Word 文档。请安装 python-docx。") from exc

    document = docx.Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _document_theme(filename: str, text: str) -> str:
    excerpt = text[:5000]
    return (
        f"从用户导入文档《{filename or '未命名文档'}》改编为可玩的 NPC 沙盒世界。\n"
        "必须抽取或改写：主线目标、支线任务、情节点、NPC、道具、道具用途、地点、胜利条件。\n"
        "如果原文是设定集或故事梗概，要补齐可跑的五步闭环和起点引导。\n\n"
        f"文档正文摘录：\n{excerpt}"
    )


def _document_fallback_world(request: WorldGenerateRequest, filename: str, text: str) -> SandboxWorldConfig:
    fallback_request = WorldGenerateRequest(
        template="short_drama_reversal",
        theme=_short_summary(filename, text),
        player_name=request.player_name,
        world_name=request.world_name or f"{Path(filename or '导入文档').stem} 改编世界",
    )
    config = generate_world_config(fallback_request)
    config.world_id = "document_import"
    config.metadata = {
        **(config.metadata or {}),
        "generated_by": "document_import_fallback",
        "source_filename": filename,
        "source_excerpt": text[:1200],
        "imported_document": True,
    }
    return config


def _short_summary(filename: str, text: str) -> str:
    lines = [line for line in text.splitlines() if line.strip()]
    payload = {"filename": filename, "excerpt": "\n".join(lines[:20])[:1500]}
    return json.dumps(payload, ensure_ascii=False)


def _looks_like_script_case(text: str) -> bool:
    legacy_case = any(token in text for token in ["剧本杀", "案件真相", "禁止提前泄露"]) and all(
        token in text for token in ["公共背景", "角色", "线索"]
    )
    asset_document = all(
        token in text
        for token in [
            "世界观",
            "人物",
            "场景",
            "时间线",
            "线索",
        ]
    ) and any(token in text for token in ["constraints.json", "约束规则", "tasks.json", "任务目标"])
    return legacy_case or asset_document
